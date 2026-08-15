"""简历解析工具：.pdf / .docx / .txt → 全文文本 + 规则提取基础信息。

设计说明：
- 文本提取支持三种主流格式（pdf 用 pypdf，docx 用 python-docx，txt 直接读）；
- 基础信息（姓名/联系方式/学历/年限/专业）用正则稳定提取——可复现、零幻觉；
  语义层面的匹配交给 Agent 的 LLM 复核，两者分工明确。
"""

from __future__ import annotations

import re
import sys
from datetime import datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

# ---------------------------------------------------------------------------
# 提取规则（正则）
# ---------------------------------------------------------------------------
PHONE_RE = re.compile(r"1[3-9]\d{9}")
EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
EDU_RE = re.compile(r"(博士|硕士|研究生|本科|大专|专科)")
YEARS_RE = re.compile(r"(\d{1,2})\s*年\s*(?:以上)?\s*(?:相关)?(?:工作)?经验")
# 工作经历时间线：如 "2024.07 - 至今" / "2020.03-2023.06"（只取起止日期，供年限推算）
WORK_RANGE_RE = re.compile(r"(20\d{2})\s*[.年]\s*\d{1,2}\s*[-–—~至到]\s*(至今|现在|今|20\d{2})")
# 教育背景行识别：这些行里的时间区间是就读时间，不算工作年限
EDU_LINE_RE = re.compile(r"大学|学院|学校|教育背景|学历")
# 教育区间的结束年份（毕业年）：如 "2021.09 - 2025.06"
GRAD_END_RE = re.compile(r"(20\d{2})\s*[.年]\s*\d{1,2}\s*[-–—~至到]\s*(20\d{2})")
MAJOR_RE = re.compile(r"(计算机科学与技术|软件工程|人工智能|数据科学|电子信息|通信工程|自动化|数学与应用数学|信息管理与信息系统)")

# 姓名启发式：第一行较短、不包含简历通用词（仅作兜底，LLM 复核会修正）
NAME_BLACKLIST = ("简历", "求职", "个人", "resume", "姓名", "男", "女")


def _detect_years(text: str) -> float | None:
    """提取工作年限，三级策略：

    1) 简历明确写「X 年经验」→ 直接取；
    2) 否则从工作/实习经历的时间线推算（最早一段经历的起始年到当前年份，
       教育背景行会被跳过）——应届生常用「2024.07 - 至今」写法，必须覆盖；
    3) 再否则用教育区间结束年（毕业年）兜底。
    返回 None 表示无法识别（交由 LLM 复核判断）。
    """
    m = YEARS_RE.search(text)
    if m:
        return float(m.group(1))

    starts: list[int] = []
    for line in text.splitlines():
        if EDU_LINE_RE.search(line):
            continue  # 教育背景行不算工作经历
        mm = WORK_RANGE_RE.search(line)
        if mm:
            starts.append(int(mm.group(1)))
    if starts:
        return max(datetime.now().year - min(starts), 0)

    for line in text.splitlines():
        if EDU_LINE_RE.search(line):
            mm = GRAD_END_RE.search(line)
            if mm:
                return max(datetime.now().year - int(mm.group(2)), 0)
    return None


def extract_text(path: str | Path) -> str:
    """按扩展名提取简历全文。"""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在：{path}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader  # 局部导入：仅解析 pdf 时需要
        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(pages).strip()
    if suffix == ".docx":
        import docx  # 局部导入：仅解析 docx 时需要
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:  # docx 表格里的经历也要读到
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts).strip()
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="ignore").strip()
    raise ValueError(f"不支持的简历格式：{suffix}（支持 .pdf / .docx / .txt）")


def extract_basics(text: str) -> dict:
    """规则提取基础信息。返回字段缺失时为 None，不硬猜。"""
    basics: dict = {}

    # 姓名：取首个非空且未被黑名单命中的短行（<=4 个字符）
    for line in text.splitlines():
        line = line.strip()
        if not line or len(line) > 6:
            continue
        if any(word in line for word in NAME_BLACKLIST):
            continue
        if PHONE_RE.search(line) or EMAIL_RE.search(line):
            continue
        basics["name"] = line
        break
    basics.setdefault("name", None)

    phone = PHONE_RE.search(text)
    email = EMAIL_RE.search(text)
    basics["phone"] = phone.group(0) if phone else None
    basics["email"] = email.group(0) if email else None

    edu = EDU_RE.search(text)
    basics["education"] = edu.group(1) if edu else None

    basics["years"] = _detect_years(text)

    major = MAJOR_RE.search(text)
    basics["major"] = major.group(1) if major else None
    return basics


def parse_resume(path: str | Path) -> dict:
    """完整解析：全文 + 基础信息（CLI 与模块共用入口）。"""
    text = extract_text(path)
    return {"file": str(Path(path).resolve()), "text": text, "basics": extract_basics(text)}


if __name__ == "__main__":
    import argparse
    import json

    parser = argparse.ArgumentParser(description="解析简历文件（.pdf/.docx/.txt）")
    parser.add_argument("--file", required=True, help="简历文件路径")
    args = parser.parse_args()
    try:
        print(json.dumps(parse_resume(args.file), ensure_ascii=False))
    except Exception as exc:
        print(json.dumps({"error": str(exc)}, ensure_ascii=False))
        sys.exit(1)
